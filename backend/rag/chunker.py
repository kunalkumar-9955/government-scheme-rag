"""
rag/chunker.py — Document parsing, cleaning, section detection, and chunking pipeline.
Supports: PDF, DOCX, HTML, and TXT with page number and rich metadata preservation.
"""
import io
import os
import re
import unicodedata
import logging
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
# 1. Text Cleaning & Normalization
# ─────────────────────────────────────────────────────────────
class DocumentCleaner:
    """
    Cleans raw extracted text from government documents:
    - Normalizes unicode whitespace and special characters
    - Fixes hyphenated line-breaks (e.g. 'bene-\n ficiary' -> 'beneficiary')
    - Strips recurring header/footer artifacts and page number stamps
    - Normalizes excessive blank lines and whitespace
    """

    @staticmethod
    def clean(text: str) -> str:
        if not text:
            return ""

        # 1. Normalize unicode (NFKC)
        text = unicodedata.normalize("NFKC", text)

        # 2. Replace non-breaking spaces, zero-width chars
        text = text.replace("\xa0", " ").replace("\u200b", "").replace("\ufeff", "")

        # 3. Normalize smart quotes and dashes
        text = re.sub(r"[\u2018\u2019\u201b]", "'", text)
        text = re.sub(r"[\u201c\u201d\u201f]", '"', text)
        text = re.sub(r"[\u2013\u2014]", "-", text)

        # 4. Fix hyphenated word breaks across newlines (e.g. "disburse-\n ment" -> "disbursement")
        text = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1\2", text)

        # 5. Remove standalone page number indicators (e.g. "Page 1 of 12", "- 1 -", "Page | 1")
        text = re.sub(r"(?i)\bPage\s*[:|]?\s*\d+\s*(of\s*\d+)?\b", "", text)
        text = re.sub(r"^\s*[-—]\s*\d+\s*[-—]\s*$", "", text, flags=re.MULTILINE)

        # 6. Normalize multiple blank lines to at most 2 newlines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # 7. Normalize spaces within lines
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        cleaned = "\n".join(lines).strip()

        return cleaned


# ─────────────────────────────────────────────────────────────
# 2. HTML Structured Parser (Zero external dependency)
# ─────────────────────────────────────────────────────────────
class HTMLToTextExtractor(HTMLParser):
    """
    Extracts structured text, headings, lists, and tables from HTML content.
    Strips script, style, nav, header, and footer tags.
    """

    IGNORABLE_TAGS = {"script", "style", "nav", "footer", "header", "noscript", "svg"}

    def __init__(self):
        super().__init__()
        self.ignore_depth = 0
        self.output_blocks = []
        self.current_block = []
        self.current_heading = ""
        self.sections = []  # List of {"title": str, "content": str}
        self.in_table = False
        self.table_rows = []
        self.current_row = []

    def handle_starttag(self, tag, attrs):
        tag_lower = tag.lower()
        if tag_lower in self.IGNORABLE_TAGS:
            self.ignore_depth += 1
            return

        if self.ignore_depth > 0:
            return

        if tag_lower in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush_block()
            self.current_heading = ""
        elif tag_lower == "table":
            self._flush_block()
            self.in_table = True
            self.table_rows = []
        elif tag_lower == "tr":
            self.current_row = []
        elif tag_lower in ("p", "div", "li", "br"):
            if not self.in_table and self.current_block:
                self._flush_block()

    def handle_endtag(self, tag):
        tag_lower = tag.lower()
        if tag_lower in self.IGNORABLE_TAGS:
            self.ignore_depth = max(0, self.ignore_depth - 1)
            return

        if self.ignore_depth > 0:
            return

        if tag_lower in ("h1", "h2", "h3", "h4", "h5", "h6"):
            heading_text = "".join(self.current_block).strip()
            self.current_block = []
            if heading_text:
                self.output_blocks.append(f"\n## {heading_text}\n")
                self.current_heading = heading_text
        elif tag_lower == "tr":
            if self.current_row:
                self.table_rows.append(self.current_row)
                self.current_row = []
        elif tag_lower == "table":
            self.in_table = False
            self._format_markdown_table()
        elif tag_lower in ("p", "div", "li"):
            self._flush_block()

    def handle_data(self, data):
        if self.ignore_depth > 0:
            return
        cleaned = data.strip()
        if cleaned:
            if self.in_table:
                self.current_row.append(cleaned)
            else:
                self.current_block.append(data)

    def _flush_block(self):
        if self.current_block:
            text = " ".join("".join(self.current_block).split())
            if text:
                self.output_blocks.append(text)
            self.current_block = []

    def _format_markdown_table(self):
        if not self.table_rows:
            return
        num_cols = max(len(r) for r in self.table_rows)
        md_lines = []
        # Header
        first_row = self.table_rows[0] + [""] * (num_cols - len(self.table_rows[0]))
        md_lines.append("| " + " | ".join(first_row) + " |")
        md_lines.append("| " + " | ".join(["---"] * num_cols) + " |")
        # Body
        for row in self.table_rows[1:]:
            padded = row + [""] * (num_cols - len(row))
            md_lines.append("| " + " | ".join(padded) + " |")
        self.output_blocks.append("\n" + "\n".join(md_lines) + "\n")

    def get_text(self) -> str:
        self._flush_block()
        return "\n\n".join(self.output_blocks)


# ─────────────────────────────────────────────────────────────
# 3. Section Detector & Chunk Classifier
# ─────────────────────────────────────────────────────────────
class SectionDetector:
    """
    Detects section headers, structure, and semantic classifications for government schemes.
    """

    ELIGIBILITY_PATTERNS = [
        r"(?i)\b(eligib\w*|who can apply|beneficiar\w* criteria|eligibility criteria|criteria for eligibility|target group)\b",
    ]
    BENEFITS_PATTERNS = [
        r"(?i)\b(benefits?|assistance|financial assistance|quantum of assistance|subsid\w*|grant|pension|coverage|entitlement)\b",
    ]
    PROCEDURE_PATTERNS = [
        r"(?i)\b(how to apply|application (procedure|process|steps)|submission|registration process|where to apply)\b",
    ]
    DOCUMENTS_PATTERNS = [
        r"(?i)\b(documents? required|required documents?|list of documents?|documentation|mandatory documents?|certificates?)\b",
    ]

    HEADING_REGEX = re.compile(
        r"^(?:"
        r"(?:\d+\.?\d*\.?\s+[A-Z][\w\s\-(),/]{2,100})|"  # "1. Introduction", "2.1 Eligibility"
        r"(?:(?:Section|Chapter|Part|Clause)\s+\w+[:\-]?\s*.*)|"  # "Section 3 - Benefits"
        r"(?:[A-Z\s]{4,80}:?$)"  # "ELIGIBILITY CRITERIA"
        r")$",
        re.MULTILINE,
    )

    @classmethod
    def classify(cls, text: str, heading: str = "") -> str:
        """Classify chunk type based on heading and content keywords."""
        combined = f"{heading} {text[:500]}".lower()

        for pat in cls.ELIGIBILITY_PATTERNS:
            if re.search(pat, combined):
                return "ELIGIBILITY"

        for pat in cls.BENEFITS_PATTERNS:
            if re.search(pat, combined):
                return "BENEFITS"

        for pat in cls.PROCEDURE_PATTERNS:
            if re.search(pat, combined):
                return "PROCEDURE"

        for pat in cls.DOCUMENTS_PATTERNS:
            if re.search(pat, combined):
                return "DOCUMENTS"

        if text.strip().startswith("|") and "---" in text:
            return "TABLE"

        if len(text.strip().splitlines()) <= 2 and len(text.strip()) < 120 and text.strip().isupper():
            return "HEADING"

        return "TEXT"

    @classmethod
    def detect_headings(cls, text: str) -> list[tuple[int, str]]:
        """Find all heading start positions and titles in text."""
        headings = []
        for match in cls.HEADING_REGEX.finditer(text):
            title = match.group(0).strip(":# ").strip()
            if len(title) >= 3:
                headings.append((match.start(), title))
        return headings


# ─────────────────────────────────────────────────────────────
# 4. Multi-format Document Parser
# ─────────────────────────────────────────────────────────────
class DocumentParser:
    """
    Unified multi-format parser for PDF, DOCX, HTML, and TXT files.
    Preserves page numbers, tables, and section structure.
    """

    def parse(self, file_source: Any, file_name: str, mime_type: str = "") -> dict:
        """
        Parse document from file path, bytes, or file-like object.
        Returns:
            {
                "full_text": str,
                "pages": [
                    {
                        "page_num": int,
                        "text": str,
                        "tables": [{"markdown": str}]
                    }
                ],
                "metadata": {
                    "page_count": int,
                    "file_name": str,
                    "mime_type": str
                }
            }
        """
        content_bytes = self._read_bytes(file_source)
        file_name_lower = (file_name or "").lower()
        mime_type_lower = (mime_type or "").lower()

        if file_name_lower.endswith(".pdf") or "pdf" in mime_type_lower:
            parsed = self._parse_pdf(content_bytes)
        elif file_name_lower.endswith((".docx", ".doc")):
            parsed = self._parse_docx(content_bytes)
        elif file_name_lower.endswith((".html", ".htm")) or "html" in mime_type_lower:
            parsed = self._parse_html(content_bytes)
        else:
            parsed = self._parse_txt(content_bytes)

        parsed["metadata"] = {
            "page_count": len(parsed.get("pages", [])),
            "file_name": file_name,
            "mime_type": mime_type,
        }
        return parsed

    def _read_bytes(self, source: Any) -> bytes:
        if isinstance(source, bytes):
            return source
        elif hasattr(source, "read"):
            source.seek(0)
            data = source.read()
            source.seek(0)
            return data
        elif isinstance(source, (str, Path)):
            p = Path(source)
            if p.is_file():
                return p.read_bytes()
            # If not direct file path, check MEDIA_ROOT
            from django.conf import settings
            media_p = Path(getattr(settings, "MEDIA_ROOT", "")) / source
            if media_p.is_file():
                return media_p.read_bytes()
            raise FileNotFoundError(f"File not found: {source}")
        else:
            raise ValueError(f"Unsupported file source type: {type(source)}")

    def _parse_pdf(self, content: bytes) -> dict:
        """Parse PDF preserving per-page text, page numbers, and tables."""
        pages = []
        full_text_parts = []

        # Try PyMuPDF (fitz) first
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            for page_idx, page in enumerate(doc, start=1):
                raw_text = page.get_text("text")
                cleaned_text = DocumentCleaner.clean(raw_text)

                tables = []
                try:
                    tab_finder = page.find_tables()
                    for tab in tab_finder.tables:
                        table_df = tab.to_pandas()
                        tables.append({
                            "markdown": table_df.to_markdown(index=False),
                            "rows": len(table_df),
                            "cols": len(table_df.columns),
                        })
                except Exception:
                    pass

                pages.append({
                    "page_num": page_idx,
                    "text": cleaned_text,
                    "tables": tables,
                })
                if cleaned_text:
                    full_text_parts.append(cleaned_text)
            doc.close()

        except ImportError:
            # Fallback to pypdf
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            for page_idx, page in enumerate(reader.pages, start=1):
                raw_text = page.extract_text() or ""
                cleaned_text = DocumentCleaner.clean(raw_text)
                pages.append({
                    "page_num": page_idx,
                    "text": cleaned_text,
                    "tables": [],
                })
                if cleaned_text:
                    full_text_parts.append(cleaned_text)

        return {
            "full_text": "\n\n".join(full_text_parts),
            "pages": pages,
        }

    def _parse_docx(self, content: bytes) -> dict:
        """Parse DOCX extracting paragraphs, heading levels, and tables."""
        from docx import Document
        doc = Document(io.BytesIO(content))

        full_text_parts = []
        tables = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                full_text_parts.append(f"\n## {text}\n")
            else:
                full_text_parts.append(text)

        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
                table_md = f"{header}\n{separator}\n{body}"
                tables.append({"markdown": table_md})
                full_text_parts.append(table_md)

        full_text = DocumentCleaner.clean("\n\n".join(full_text_parts))
        return {
            "full_text": full_text,
            "pages": [{"page_num": 1, "text": full_text, "tables": tables}],
        }

    def _parse_html(self, content: bytes) -> dict:
        """Parse HTML/Web content using clean semantic extractor."""
        html_str = content.decode("utf-8", errors="replace")
        extractor = HTMLToTextExtractor()
        extractor.feed(html_str)
        cleaned_text = DocumentCleaner.clean(extractor.get_text())

        return {
            "full_text": cleaned_text,
            "pages": [{"page_num": 1, "text": cleaned_text, "tables": []}],
        }

    def _parse_txt(self, content: bytes) -> dict:
        """Parse plain text."""
        text = content.decode("utf-8", errors="replace")
        cleaned_text = DocumentCleaner.clean(text)
        return {
            "full_text": cleaned_text,
            "pages": [{"page_num": 1, "text": cleaned_text, "tables": []}],
        }


# ─────────────────────────────────────────────────────────────
# 5. Document Chunker with Rich Metadata Preservation
# ─────────────────────────────────────────────────────────────
class DocumentChunker:
    """
    Multi-strategy chunker that preserves page numbers and all required metadata:
    - document_id
    - scheme_id
    - scheme_name
    - ministry
    - department
    - state
    - category
    - section
    - page_number
    - source_url
    - document_version
    """

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64, strategy: str = "recursive"):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy

    def chunk(self, parsed_content: dict, metadata: dict) -> list[dict]:
        """
        Chunk parsed document preserving page_number and all metadata on each chunk.
        """
        chunks = []
        current_section = metadata.get("section") or "General"

        # Chunk page-by-page to accurately preserve page_number
        for page in parsed_content.get("pages", []):
            page_num = page.get("page_num", 1)
            page_text = page.get("text", "")

            # 1. Standalone Table chunks
            for table in page.get("tables", []):
                if table.get("markdown"):
                    table_chunk = self._make_chunk(
                        content=table["markdown"],
                        chunk_type="TABLE",
                        page_number=page_num,
                        section_title=current_section,
                        metadata=metadata,
                    )
                    chunks.append(table_chunk)

            if not page_text.strip():
                continue

            # 2. Text chunks for this page
            if self.strategy == "semantic":
                page_chunks = self._semantic_chunk(page_text, page_num, current_section, metadata)
            elif self.strategy == "fixed":
                page_chunks = self._fixed_chunk(page_text, page_num, current_section, metadata)
            else:  # recursive (default)
                page_chunks = self._recursive_chunk(page_text, page_num, current_section, metadata)

            for c in page_chunks:
                # Update current_section if a heading is detected
                if c.get("section_title") and c["section_title"] != "General":
                    current_section = c["section_title"]
                chunks.append(c)

        return chunks

    def _make_chunk(
        self,
        content: str,
        chunk_type: str,
        page_number: Optional[int],
        section_title: str,
        metadata: dict,
    ) -> dict:
        """Create standard chunk dict with all 11 required metadata fields."""
        from core.utils import count_tokens

        section = section_title or metadata.get("section") or "General"

        # Explicitly preserve all 11 required metadata attributes
        preserved_metadata = {
            "document_id": str(metadata.get("document_id") or ""),
            "scheme_id": str(metadata.get("scheme_id") or "") if metadata.get("scheme_id") else None,
            "scheme_name": str(metadata.get("scheme_name") or metadata.get("title") or ""),
            "ministry": str(metadata.get("ministry") or ""),
            "department": str(metadata.get("department") or ""),
            "state": str(metadata.get("state") or ""),
            "category": str(metadata.get("category") or ""),
            "section": section,
            "page_number": page_number,
            "source_url": str(metadata.get("source_url") or ""),
            "document_version": str(metadata.get("document_version") or "1.0"),
        }

        # Include any extra custom metadata passed
        for k, v in metadata.items():
            if k not in preserved_metadata:
                preserved_metadata[k] = v

        return {
            "content": content.strip(),
            "chunk_type": chunk_type,
            "page_number": page_number,
            "section_title": section,
            "token_count": count_tokens(content),
            "char_count": len(content.strip()),
            "metadata": preserved_metadata,
        }

    def _recursive_chunk(self, text: str, page_number: int, section_title: str, metadata: dict) -> list[dict]:
        """LangChain-compatible recursive character text splitter."""
        char_size = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4

        raw_pieces = []
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=char_size,
                chunk_overlap=char_overlap,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
            raw_pieces = splitter.split_text(text)
        except ImportError:
            raw_pieces = self._fallback_split(text, char_size, char_overlap)

        chunks = []
        active_heading = section_title

        for piece in raw_pieces:
            piece_str = piece.strip()
            if not piece_str:
                continue

            # Check if piece contains a new section heading
            headings = SectionDetector.detect_headings(piece_str)
            if headings:
                active_heading = headings[0][1]

            chunk_type = SectionDetector.classify(piece_str, active_heading)

            chunks.append(self._make_chunk(
                content=piece_str,
                chunk_type=chunk_type,
                page_number=page_number,
                section_title=active_heading,
                metadata=metadata,
            ))

        return chunks

    def _semantic_chunk(self, text: str, page_number: int, section_title: str, metadata: dict) -> list[dict]:
        """Split on semantic section boundaries."""
        sections = re.split(r'\n(?=(?:[0-9]+\.?\s+[A-Z]|[A-Z\s]{5,}:?\n))', text)
        chunks = []
        active_heading = section_title

        for section in sections:
            section_str = section.strip()
            if not section_str:
                continue

            headings = SectionDetector.detect_headings(section_str)
            if headings:
                active_heading = headings[0][1]

            if len(section_str) > self.chunk_size * 5:
                # Sub-chunk large sections recursively
                sub_chunks = self._recursive_chunk(section_str, page_number, active_heading, metadata)
                chunks.extend(sub_chunks)
            else:
                chunk_type = SectionDetector.classify(section_str, active_heading)
                chunks.append(self._make_chunk(
                    content=section_str,
                    chunk_type=chunk_type,
                    page_number=page_number,
                    section_title=active_heading,
                    metadata=metadata,
                ))

        return chunks

    def _fixed_chunk(self, text: str, page_number: int, section_title: str, metadata: dict) -> list[dict]:
        """Fixed character size chunks with overlap."""
        char_size = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4
        raw_pieces = self._fallback_split(text, char_size, char_overlap)

        chunks = []
        for piece in raw_pieces:
            chunk_type = SectionDetector.classify(piece, section_title)
            chunks.append(self._make_chunk(
                content=piece,
                chunk_type=chunk_type,
                page_number=page_number,
                section_title=section_title,
                metadata=metadata,
            ))
        return chunks

    def _fallback_split(self, text: str, size: int, overlap: int) -> list[str]:
        """Fallback character chunker with overlap."""
        paragraphs = re.split(r"\n\s*\n", text)
        pieces = []
        current = ""

        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
            if len(current) + len(p) + 2 <= size:
                current = f"{current}\n\n{p}".strip()
            else:
                if current:
                    pieces.append(current)
                if len(p) > size:
                    # Slice large paragraph
                    start = 0
                    while start < len(p):
                        end = min(start + size, len(p))
                        pieces.append(p[start:end])
                        start = end - overlap if end < len(p) else end
                    current = ""
                else:
                    current = p

        if current:
            pieces.append(current)
        return pieces
