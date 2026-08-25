"""
apps/documents/tests.py — Comprehensive Unit Tests for Document Ingestion Pipeline
Covers:
- PDF extraction with page number preservation
- DOCX extraction with table & heading preservation
- HTML extraction with semantic structure
- Chunking strategies & section classification
- Metadata preservation on every chunk (all 11 fields)
- Duplicate document detection via SHA-256
- File validation & anti-malware security checks
- End-to-end ingestion task execution
"""
import io
import json
from decimal import Decimal
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APITestCase

from apps.authentication.models import CustomUser
from core.permissions import UserRole
from apps.schemes.models import GovernmentScheme, SchemeType, SchemeStatus
from apps.documents.models import GovDocument, DocumentChunk, DocumentStatus, DocumentCategory
from rag.chunker import DocumentParser, DocumentChunker, DocumentCleaner, SectionDetector
from rag.embedder import EmbeddingService
from apps.documents.tasks import process_document


def _create_sample_pdf_bytes() -> bytes:
    """Helper to generate a valid 2-page PDF in-memory using pypdf."""
    import pypdf
    writer = pypdf.PdfWriter()

    # Page 1
    page1_text = (
        "Pradhan Mantri Krishi Sinchayee Yojana (PMKSY)\n"
        "Ministry of Agriculture & Farmers Welfare\n\n"
        "1. Objectives\n"
        "The overarching vision of PMKSY is to ensure access to some means of protective irrigation to all agricultural farms in the country."
    )
    # Page 2
    page2_text = (
        "2. Eligibility Criteria\n"
        "All farmers and landholders across all states are eligible for micro-irrigation subsidies.\n\n"
        "3. Benefits\n"
        "Financial assistance of up to 55% for small and marginal farmers and 45% for other farmers for installing drip and sprinkler irrigation systems."
    )

    # Use pypdf to add pages with annotations/text stream
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)

    # Let's write PDF with real stream text objects
    output = io.BytesIO()
    writer.write(output)
    pdf_bytes = output.getvalue()

    # Inject page stream content so text extraction extracts text
    # Or create via PyPDF / minimal valid text PDF
    minimal_pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 5 0 R /Resources << /Font << /F1 7 0 R >> >> >> endobj\n"
        b"4 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 6 0 R /Resources << /Font << /F1 7 0 R >> >> >> endobj\n"
        b"5 0 obj << /Length 120 >> stream\n"
        b"BT /F1 12 Tf 50 700 Td (PMKSY Guidelines - Page 1. Objective: Provide irrigation to all farms.) Tj ET\n"
        b"endstream\nendobj\n"
        b"6 0 obj << /Length 140 >> stream\n"
        b"BT /F1 12 Tf 50 700 Td (Eligibility Criteria: All farmers. Benefits: Up to 55 percent subsidy.) Tj ET\n"
        b"endstream\nendobj\n"
        b"7 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"xref\n0 8\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
        b"0000000121 00000 n \n0000000244 00000 n \n0000000367 00000 n \n0000000538 00000 n \n"
        b"0000000729 00000 n \ntrailer << /Size 8 /Root 1 0 R >>\nstartxref\n806\n%%EOF"
    )
    return minimal_pdf


def _create_sample_docx_bytes() -> bytes:
    """Helper to generate a valid DOCX file in-memory using python-docx."""
    import docx
    doc = docx.Document()
    doc.add_heading("Pradhan Mantri Awas Yojana", level=1)
    doc.add_paragraph("PMAY provides affordable pucca housing to rural and urban poor.")
    doc.add_heading("Eligibility Criteria", level=2)
    doc.add_paragraph("Families with annual household income below Rs 3,00,000 without a pucca house.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Category"
    table.rows[0].cells[1].text = "Assistance"
    table.rows[1].cells[0].text = "Plain Area"
    table.rows[1].cells[1].text = "Rs. 1,20,000"

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


class DocumentIngestionPipelineTests(APITestCase):
    def setUp(self):
        # Users
        self.admin = CustomUser.objects.create_user(
            email="doc_admin@example.com",
            password="AdminPassword123!",
            role=UserRole.ADMIN,
            is_staff=True,
        )
        self.citizen = CustomUser.objects.create_user(
            email="doc_citizen@example.com",
            password="CitizenPassword123!",
            role=UserRole.CITIZEN,
        )

        # Sample Scheme
        self.scheme = GovernmentScheme.objects.create(
            name="Pradhan Mantri Krishi Sinchayee Yojana",
            short_title="PMKSY",
            slug="pmksy",
            description="Irrigation support for all agricultural farms",
            scheme_type=SchemeType.CENTRAL_SECTOR,
            status=SchemeStatus.ACTIVE,
        )

        self.upload_url = reverse("document-upload")
        self.list_url = reverse("document-list")

    def test_pdf_text_extraction_and_page_numbers(self):
        """Test that PDF extraction correctly extracts text and preserves page numbers."""
        pdf_bytes = _create_sample_pdf_bytes()
        parser = DocumentParser()
        parsed = parser.parse(pdf_bytes, "guidelines.pdf", "application/pdf")

        self.assertIn("pages", parsed)
        self.assertEqual(len(parsed["pages"]), 2)
        self.assertEqual(parsed["pages"][0]["page_num"], 1)
        self.assertEqual(parsed["pages"][1]["page_num"], 2)
        self.assertIn("Objective", parsed["pages"][0]["text"])
        self.assertIn("Eligibility", parsed["pages"][1]["text"])

    def test_docx_text_and_table_extraction(self):
        """Test DOCX parsing extracts structured text, headings, and tables."""
        docx_bytes = _create_sample_docx_bytes()
        parser = DocumentParser()
        parsed = parser.parse(docx_bytes, "pmay_guidelines.docx")

        self.assertIn("Pradhan Mantri Awas Yojana", parsed["full_text"])
        self.assertIn("Eligibility Criteria", parsed["full_text"])
        self.assertGreaterEqual(len(parsed["pages"]), 1)
        self.assertGreaterEqual(len(parsed["pages"][0]["tables"]), 1)
        self.assertIn("Plain Area", parsed["pages"][0]["tables"][0]["markdown"])

    def test_html_extraction_and_cleaning(self):
        """Test HTML parser strips tags and extracts clean text and sections."""
        html_content = b"""
        <!DOCTYPE html>
        <html>
        <head><title>Scheme Portal</title><style>.banner{color:red;}</style></head>
        <body>
            <header><nav><a href="/">Home</a></nav></header>
            <h1>National Social Assistance Programme</h1>
            <p>NSAP represents a significant step towards the fulfillment of the Directive Principles.</p>
            <h2>Eligibility Criteria</h2>
            <p>Beneficiary must be aged 60 years or above and living below poverty line.</p>
            <table>
                <tr><th>Component</th><th>Monthly Pension</th></tr>
                <tr><td>IGNOAPS</td><td>Rs. 500</td></tr>
            </table>
            <footer>Copyright 2026 Government of India</footer>
            <script>console.log("analytics");</script>
        </body>
        </html>
        """
        parser = DocumentParser()
        parsed = parser.parse(html_content, "nsap.html", "text/html")

        # Must not contain script, style, nav or footer content
        self.assertNotIn("console.log", parsed["full_text"])
        self.assertNotIn("Copyright 2026", parsed["full_text"])
        self.assertIn("National Social Assistance Programme", parsed["full_text"])
        self.assertIn("Eligibility Criteria", parsed["full_text"])
        self.assertIn("IGNOAPS", parsed["full_text"])

    def test_chunking_and_metadata_preservation(self):
        """Test that every chunk preserves all 11 required metadata fields and page numbers."""
        pdf_bytes = _create_sample_pdf_bytes()
        parser = DocumentParser()
        parsed = parser.parse(pdf_bytes, "pmksy.pdf")

        chunker = DocumentChunker(chunk_size=128, chunk_overlap=16, strategy="recursive")
        metadata = {
            "document_id": "test-doc-uuid-1234",
            "scheme_id": str(self.scheme.id),
            "scheme_name": self.scheme.name,
            "ministry": "Ministry of Agriculture",
            "department": "Department of Agriculture & Cooperation",
            "state": "All India",
            "category": "AGRICULTURE",
            "section": "General",
            "source_url": "https://pmksy.gov.in/guidelines.pdf",
            "document_version": "2.1",
        }

        chunks = chunker.chunk(parsed, metadata=metadata)
        self.assertGreaterEqual(len(chunks), 2)

        # Verify EVERY chunk contains the 11 metadata fields
        required_keys = [
            "document_id", "scheme_id", "scheme_name", "ministry", "department",
            "state", "category", "section", "page_number", "source_url", "document_version"
        ]

        for chunk in chunks:
            self.assertIn("content", chunk)
            self.assertIn("chunk_type", chunk)
            self.assertIn("page_number", chunk)
            self.assertIsNotNone(chunk["page_number"])
            meta = chunk["metadata"]
            for key in required_keys:
                self.assertIn(key, meta, f"Missing required metadata key: {key}")
            self.assertEqual(meta["document_id"], "test-doc-uuid-1234")
            self.assertEqual(meta["scheme_name"], self.scheme.name)
            self.assertEqual(meta["document_version"], "2.1")

    def test_section_detector_and_chunk_classification(self):
        """Test section detection identifies ELIGIBILITY, BENEFITS, and PROCEDURE."""
        eligibility_text = "Eligibility Criteria: The applicant must be a resident citizen of India."
        benefits_text = "Quantum of Financial Assistance: Rs. 6000 per year paid in 3 installments."
        procedure_text = "How to Apply: Visit the official portal and submit the registration form."

        self.assertEqual(SectionDetector.classify(eligibility_text), "ELIGIBILITY")
        self.assertEqual(SectionDetector.classify(benefits_text), "BENEFITS")
        self.assertEqual(SectionDetector.classify(procedure_text), "PROCEDURE")

    def test_duplicate_document_detection_via_hash(self):
        """Test uploading the exact same document twice returns 409 DUPLICATE_DOCUMENT."""
        self.client.force_authenticate(user=self.admin)
        pdf_bytes = _create_sample_pdf_bytes()

        file1 = SimpleUploadedFile("pmksy_v1.pdf", pdf_bytes, content_type="application/pdf")
        data1 = {
            "file": file1,
            "title": "PMKSY Operational Guidelines",
            "ministry": "Ministry of Agriculture",
            "category": "AGRICULTURE",
        }
        res1 = self.client.post(self.upload_url, data1, format="multipart")
        self.assertEqual(res1.status_code, status.HTTP_201_CREATED)
        doc_id = res1.data["data"]["id"]

        # Attempt to upload identical file
        file2 = SimpleUploadedFile("pmksy_copy.pdf", pdf_bytes, content_type="application/pdf")
        data2 = {
            "file": file2,
            "title": "PMKSY Duplicate Guidelines",
            "ministry": "Ministry of Agriculture",
            "category": "AGRICULTURE",
        }
        res2 = self.client.post(self.upload_url, data2, format="multipart")
        self.assertEqual(res2.status_code, status.HTTP_409_CONFLICT)
        self.assertEqual(res2.data["error"]["code"], "DUPLICATE_DOCUMENT")
        self.assertEqual(res2.data["error"]["existing_document_id"], str(doc_id))

    def test_malicious_and_invalid_file_rejection(self):
        """Test rejecting non-allowed extensions and fake binary headers."""
        self.client.force_authenticate(user=self.admin)

        # 1. Disallowed extension (.exe)
        bad_file = SimpleUploadedFile("malware.exe", b"MZ\x90\x00\x03\x00", content_type="application/octet-stream")
        res = self.client.post(self.upload_url, {"file": bad_file, "title": "Malware Executable"}, format="multipart")
        self.assertEqual(res.status_code, status.HTTP_400_BAD_REQUEST)

        # 2. Fake PDF with Windows executable MZ magic header
        fake_pdf = SimpleUploadedFile("fake.pdf", b"MZ\x90\x00\x03\x00FakePEHeaderInsidePDF", content_type="application/pdf")
        res_fake = self.client.post(self.upload_url, {"file": fake_pdf, "title": "Fake PDF"}, format="multipart")
        self.assertEqual(res_fake.status_code, status.HTTP_400_BAD_REQUEST)

    def test_end_to_end_ingestion_pipeline_task(self):
        """Test end-to-end async/sync execution of process_document."""
        self.client.force_authenticate(user=self.admin)
        docx_bytes = _create_sample_docx_bytes()
        file = SimpleUploadedFile("pmay.docx", docx_bytes, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        res = self.client.post(
            self.upload_url,
            {
                "file": file,
                "title": "PMAY Official Guidelines",
                "scheme": str(self.scheme.id),
                "ministry": "Ministry of Rural Development",
                "department": "Rural Housing",
                "category": "HOUSING",
                "document_version": "3.0",
                "state": "National",
            },
            format="multipart",
        )
        self.assertEqual(res.status_code, status.HTTP_201_CREATED)
        doc_id = res.data["data"]["id"]

        # Run process_document pipeline
        process_document(doc_id)

        doc = GovDocument.objects.get(id=doc_id)
        self.assertEqual(doc.status, DocumentStatus.COMPLETED)
        self.assertGreater(doc.total_chunks, 0)
        self.assertIsNotNone(doc.processed_at)

        # Verify created chunks in DB
        chunks = DocumentChunk.objects.filter(document=doc)
        self.assertGreaterEqual(chunks.count(), 1)
        first_chunk = chunks.first()
        self.assertIn("document_id", first_chunk.metadata)
        self.assertEqual(first_chunk.metadata["scheme_id"], str(self.scheme.id))
        self.assertEqual(first_chunk.metadata["document_version"], "3.0")
        self.assertEqual(first_chunk.metadata["state"], "National")
        self.assertIn("embedding", first_chunk.metadata)
        self.assertEqual(len(first_chunk.metadata["embedding"]), 768)
